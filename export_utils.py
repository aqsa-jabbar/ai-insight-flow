"""
export_utils.py
----------------
Converts a plain-text report into downloadable DOCX and PDF files, entirely
in-memory (using BytesIO) so nothing is written to disk - important for
Hugging Face Spaces, where the filesystem is ephemeral anyway.

Both functions do simple Markdown-ish parsing (#, ##, ###, and "- " bullets)
so headings in the LLM's output turn into real headings in the exported
file, instead of just plain paragraphs.

NOTE ON THE PDF WRAPPER:
FPDF's own multi_cell() has a fragile internal line-break algorithm that can
throw FPDFException("Not enough horizontal space to render a single
character") on certain inputs (long dash runs, table-like text, etc.),
even after pre-chunking long words. To avoid this entirely, we measure
text with the font's real character widths ourselves (get_string_width)
and print pre-wrapped, guaranteed-to-fit lines one at a time with cell().
This bypasses FPDF's internal wrapper completely, so the exception can't
happen.
"""

from io import BytesIO

from docx import Document
from fpdf import FPDF


def generate_docx(title: str, report_text: str) -> bytes:
    """
    Builds a .docx file from the report text and returns it as raw bytes,
    ready to be handed to st.download_button.
    """
    doc = Document()
    doc.add_heading(title, level=1)

    for raw_line in report_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_pdf(title: str, report_text: str) -> bytes:
    """
    Builds a .pdf file from the report text and returns it as raw bytes,
    ready to be handed to st.download_button.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Helvetica", "B", 16)
    _print_wrapped(pdf, _safe_text(title), line_height=10)
    pdf.ln(4)

    pdf.set_font("Helvetica", size=11)
    for raw_line in report_text.split("\n"):
        line = raw_line.strip()
        if not line:
            pdf.ln(4)
            continue

        if line.startswith(("#", "##", "###")):
            heading = line.lstrip("#").strip()
            pdf.set_font("Helvetica", "B", 13)
            _print_wrapped(pdf, _safe_text(heading), line_height=8)
            pdf.set_font("Helvetica", size=11)
        elif line.startswith(("- ", "* ")):
            _print_wrapped(pdf, _safe_text(f"\u2022 {line[2:]}"), line_height=7)
        else:
            _print_wrapped(pdf, _safe_text(line), line_height=7)

    # fpdf2's output() returns a bytearray - convert to plain bytes
    return bytes(pdf.output(dest="S"))


def _safe_text(text: str) -> str:
    """
    FPDF's built-in Helvetica font only supports latin-1 characters.
    This replaces anything it can't render (e.g. smart quotes, emoji)
    instead of crashing the export.
    """
    return text.encode("latin-1", "replace").decode("latin-1")


def _wrap_line(pdf: FPDF, text: str, max_width: float) -> list:
    """
    Wraps a single line of text into a list of lines that each measurably
    fit within max_width, using the PDF's actual current font metrics
    (pdf.get_string_width). Any single "word" wider than max_width on its
    own (e.g. a long URL, or a long run of dashes/underscores with no
    spaces) is force-broken character by character so nothing ever
    overflows the page width.
    """
    if not text:
        return [""]

    words = text.split(" ")
    lines = []
    current = ""

    for word in words:
        # Force-break any single word that alone is wider than the page.
        while pdf.get_string_width(word) > max_width:
            cut = len(word)
            while cut > 1 and pdf.get_string_width(word[:cut]) > max_width:
                cut -= 1
            if current:
                lines.append(current)
                current = ""
            lines.append(word[:cut])
            word = word[cut:]

        candidate = f"{current} {word}".strip() if current else word
        if current and pdf.get_string_width(candidate) > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate

    if current:
        lines.append(current)

    return lines


def _print_wrapped(pdf: FPDF, text: str, line_height: float) -> None:
    """
    Wraps `text` to fit the page width and prints it line by line using
    cell() instead of multi_cell(). Because every line handed to cell()
    is already guaranteed to fit, FPDF's own (fragile) line-break logic
    never runs, so it can't throw FPDFException.
    """
    max_width = pdf.w - pdf.l_margin - pdf.r_margin
    for line in _wrap_line(pdf, text, max_width):
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, line_height, line)
        pdf.ln(line_height)